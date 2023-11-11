from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from eng_module import beams
from eng_module import plots
from PyNite import FEModel3D
import math
from typing import Optional
from io import BytesIO
from matplotlib.figure import Figure

def create_beam_report (beam_model=FEModel3D,
                        output_filename=str,
                        project_name=str,
                        designer=str,
                        force_units=str,
                        moment_units=str,
                        load_combo=str,
                    ) -> None:
    """
     writes a new .docx file to disk populated with all of the information 
     and formatting that we want for our analysis report.
    """
    doc = Document("eng_module/templates/sp_template_empty.docx")
    p1 = doc.add_paragraph(f"Calculation results for {load_combo}", style="Title")
    p2 = doc.add_paragraph("Shear: ", style="Heading 1")
    p2_run = p2.add_run("Plot of results")
    envelope_shear = plots.plot_results(beam_model, "shear", 
                 "Fy", "N", 
                 load_combo)
    my_plot_data = BytesIO() # Currently empty
    envelope_shear.savefig(my_plot_data)
    p2_run.add_picture(my_plot_data) 
    p3 = doc.add_paragraph("Moment: ", style="Heading 1")
    p3_run = p3.add_run("Plot of results")
    envelope_moment = plots.plot_results(beam_model, "moment", 
                 "Mz", "Nmm", 
                 load_combo)
    my_plot_data2 = BytesIO() # Currently empty
    envelope_moment.savefig(my_plot_data2)
    p3_run.add_picture(my_plot_data2) 
    doc.save(output_filename)
    return

